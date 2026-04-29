import os
import io
import json
import zipfile
from flask import Flask, request, send_file, send_from_directory, session, redirect, url_for
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.colors import HexColor
from docx import Document
import pandas as pd

app = Flask(__name__, static_folder='.', static_url_path='')
app.secret_key = 'super_secret_key_change_this'

UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

OBJECTS_FILE = 'objects.json'
ADMIN_LOGIN = 'admin'
ADMIN_PASSWORD = 'admin123'

# Регистрация шрифта ГОСТ
try:
    pdfmetrics.registerFont(TTFont('GostTypeB', 'GOST_A.ttf'))
    FONT_NAME = 'GostTypeB'
except Exception as e:
    print(f"ВНИМАНИЕ: Шрифт GOST_A.ttf не найден. Будет использован стандартный шрифт. Ошибка: {e}")
    FONT_NAME = 'Helvetica'

def load_objects():
    if not os.path.exists(OBJECTS_FILE):
        default_objects = [
            {"egrkn": "441410069670005", "name": "Церковь Воскресения Словущего", "address": "Костромская область, г. Кострома, ул. Симановского, д. 5", "history": "1685 г.", "fullName": "Церковь Воскресения Словущего\n1685 г."},
            {"egrkn": "441711139670005", "name": "Торговые ряды", "address": "Костромская область, г. Кострома, ул. Симановского, д. 1", "history": "1789-1791 гг.", "fullName": "Торговые ряды\n1789-1791 гг."}
        ]
        save_objects(default_objects)
        return default_objects
    with open(OBJECTS_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)

def save_objects(data):
    with open(OBJECTS_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# Создаем файл objects.json при запуске
if not os.path.exists(OBJECTS_FILE):
    load_objects()

@app.route('/')
def index():
    return send_from_directory('.', 'index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        if username == ADMIN_LOGIN and password == ADMIN_PASSWORD:
            session['logged_in'] = True
            return redirect(url_for('admin'))
        else:
            return 'Неверный логин или пароль', 401
    return send_from_directory('.', 'login.html')

@app.route('/logout')
def logout():
    session.pop('logged_in', None)
    return redirect(url_for('index'))

@app.route('/admin')
def admin():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    return send_from_directory('.', 'admin.html')

@app.route('/api/objects')
def api_objects():
    return send_file(OBJECTS_FILE)

@app.route('/api/upload-excel', methods=['POST'])
def upload_excel():
    if not session.get('logged_in'):
        return {'error': 'Unauthorized'}, 401
    if 'file' not in request.files:
        return {'error': 'No file'}, 400
    file = request.files['file']
    if file.filename == '':
        return {'error': 'No selected file'}, 400
    try:
        df = pd.read_excel(file)
        data = df.to_dict(orient='records')
        save_objects(data)
        return {'success': True, 'message': f'Загружено {len(data)} объектов'}
    except Exception as e:
        return {'error': str(e)}, 500

@app.route('/generate', methods=['POST'])
def generate_project():
    try:
        data = request.form.to_dict()
        
        # Обработка файлов
        scheme_path = None
        if 'installationScheme' in request.files and request.files['installationScheme'].filename:
            scheme_path = os.path.join(app.config['UPLOAD_FOLDER'], 'scheme.png')
            request.files['installationScheme'].save(scheme_path)
            
        photo_path = None
        if 'colorPhoto' in request.files and request.files['colorPhoto'].filename:
            photo_path = os.path.join(app.config['UPLOAD_FOLDER'], 'photo.png')
            request.files['colorPhoto'].save(photo_path)

        # --- ГЕНЕРАЦИЯ PDF ---
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=A4)
        width, height = A4

        def add_footer(canvas_obj):
            canvas_obj.setFont(FONT_NAME, 8)
            canvas_obj.drawCentredString(width / 2, 30, "Выполнено в «Цифровом конструкторе ПИН»")

        c.setFont(FONT_NAME, 14)
        c.drawCentredString(width / 2, height - 50, "ПРОЕКТ УСТАНОВКИ И СОДЕРЖАНИЯ")
        c.drawCentredString(width / 2, height - 70, "ИНФОРМАЦИОННЫХ НАДПИСЕЙ")
        
        c.setFont(FONT_NAME, 12)
        y_pos = height - 120
        c.drawString(50, y_pos, f"Объект: {data.get('objectName', '')}")
        y_pos -= 20
        c.drawString(50, y_pos, f"Адрес: {data.get('location', '')}")
        y_pos -= 20
        c.drawString(50, y_pos, f"Рег. номер ЕГРОКН: {data.get('egrknNumber', '')}")
        y_pos -= 40
        c.drawString(50, y_pos, f"Разработчик: {data.get('developer', '')}")
        y_pos -= 20
        c.drawString(50, y_pos, f"Место разработки: {data.get('devPlace', '')}")
        y_pos -= 20
        c.drawString(50, y_pos, f"Дата: {data.get('devDate', '')}")
        add_footer(c)
        c.showPage()
        c.setFont(FONT_NAME, 16)
        c.drawString(50, height - 50, "1. Общие сведения об объекте культурного наследия")
        c.setFont(FONT_NAME, 12)
        y_pos = height - 80
        
        def draw_line(label, value):
            nonlocal y_pos
            c.setFont(FONT_NAME, 10)
            c.drawString(50, y_pos, label)
            c.setFont(FONT_NAME, 12)
            c.drawString(160, y_pos, str(value or "—"))
            y_pos -= 20

        draw_line("Категория:", data.get('category'))
        draw_line("Вид объекта:", data.get('objectType'))
        draw_line("Дата охраны:", data.get('protectionDate'))
        draw_line("Пользователь:", data.get('userInfo'))
        draw_line("Рег. номер:", data.get('egrknNumber'))
        draw_line("Ранее установленные надписи:", data.get('previousSigns'))

        add_footer(c)
        c.showPage()
        c.setFont(FONT_NAME, 16)
        c.drawString(50, height - 50, "2. Эскизное предложение")
        c.setFont(FONT_NAME, 12)
        y_pos = height - 90
        
        c.drawString(50, y_pos, "Обоснование:")
        y_pos -= 20
        justification = data.get('justification', '')
        lines = [justification[i:i+80] for i in range(0, len(justification), 80)]
        for line in lines:
            c.drawString(60, y_pos, line)
            y_pos -= 15

        y_pos -= 20
        c.drawString(50, y_pos, "Текст на табличке:")
        y_pos -= 20
        plate_text = data.get('plateText', '')
        lines_plate = [plate_text[i:i+80] for i in range(0, len(plate_text), 80)]
        for line in lines_plate:
            c.drawString(60, y_pos, line)
            y_pos -= 15

        add_footer(c)
        c.showPage()
        c.setFont(FONT_NAME, 16)
        c.drawString(50, height - 50, "3. Технические характеристики и чертеж пластины")
        c.setFont(FONT_NAME, 12)
        y_pos = height - 80
        
        draw_line("Размер пластины:", data.get('plateSize'))
        draw_line("Материал стен:", data.get('wallMaterial'))
        draw_line("Масса:", f"{data.get('calcMass', '0')} кг")
        draw_line("Ветровая нагрузка:", f"{data.get('calcWindLoad', '0')} кН")
        
        y_pos -= 40
        c.setFont(FONT_NAME, 14)
        c.drawString(50, y_pos, "Чертеж пластины (Масштаб 1:5):")
        y_pos -= 60
        plate_w = 300
        plate_h = 200
        c.setFillColor(HexColor('#4a4a4a'))
        c.rect(50, y_pos - plate_h, plate_w, plate_h, fill=1)
        
        c.setFillColor(HexColor('#ffffff'))
        c.setFont(FONT_NAME, 10)
        c.drawCentredString(50 + plate_w/2, y_pos - 20, "ОБЪЕКТ КУЛЬТУРНОГО НАСЛЕДИЯ")
        c.drawCentredString(50 + plate_w/2, y_pos - 40, data.get('category', '').upper())
        c.drawCentredString(50 + plate_w/2, y_pos - 70, data.get('objectName', '').upper())
        c.drawCentredString(50 + plate_w/2, y_pos - 90, data.get('historyInfo', '').upper())
        c.drawCentredString(50 + plate_w/2, y_pos - 120, f"№ {data.get('egrknNumber', '')}")
        c.drawCentredString(50 + plate_w/2, y_pos - plate_h + 20, "ОХРАНЯЕТСЯ ГОСУДАРСТВОМ")

        add_footer(c)
        c.showPage()
        c.setFont(FONT_NAME, 16)
        c.drawString(50, height - 50, "4. Схема установки и фотофиксация")
        y_pos = height - 80
        
        if scheme_path:
            try:
                c.drawImage(scheme_path, 50, y_pos - 180, width=450, height=180)
                y_pos -= 200
                c.setFont(FONT_NAME, 10)
                c.drawString(50, y_pos + 190, "Схема установки:")
                c.setFont(FONT_NAME, 12)
            except: pass
            
        if photo_path:
            try:
                c.drawImage(photo_path, 50, y_pos - 180, width=450, height=180)
                y_pos -= 200
                c.setFont(FONT_NAME, 10)
                c.drawString(50, y_pos + 190, "Фотофиксация:")
                c.setFont(FONT_NAME, 12)
            except: pass

        add_footer(c)
        c.save()
        pdf_buffer.seek(0)

        # --- ГЕНЕРАЦИЯ DOCX ---
        doc = Document()
        doc.add_heading('Проект информационной надписи', 0)
        doc.add_paragraph(f"Объект: {data.get('objectName', '')}")
        doc.add_paragraph(f"Адрес: {data.get('location', '')}")
        doc.add_paragraph(f"Рег. номер ЕГРОКН: {data.get('egrknNumber', '')}")
        
        h1 = doc.add_heading('1. Общие сведения', level=1)
        doc.add_paragraph(f"Категория: {data.get('category', '')}")
        doc.add_paragraph(f"Вид объекта: {data.get('objectType', '')}")
        doc.add_paragraph(f"Дата охраны: {data.get('protectionDate', '')}")
        doc.add_paragraph(f"Пользователь: {data.get('userInfo', '')}")
        
        h2 = doc.add_heading('2. Эскизное предложение', level=1)
        doc.add_paragraph("Обоснование:")
        doc.add_paragraph(data.get('justification', ''))
        doc.add_paragraph("Текст на табличке:")
        doc.add_paragraph(data.get('plateText', ''))
        
        h3 = doc.add_heading('3. Технические характеристики', level=1)
        doc.add_paragraph(f"Размер пластины: {data.get('plateSize', '')}")
        doc.add_paragraph(f"Материал стен: {data.get('wallMaterial', '')}")
        doc.add_paragraph(f"Масса: {data.get('calcMass', '0')} кг")
        doc.add_paragraph(f"Ветровая нагрузка: {data.get('calcWindLoad', '0')} кН")
        doc.add_paragraph(f"Нагрузка на срез: {data.get('calcGravityLoad', '0')} кН")

        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        # --- ZIP АРХИВ ---
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            zipf.writestr('Project_PIN.pdf', pdf_buffer.read())
            zipf.writestr('Project_PIN.docx', docx_buffer.read())
        
        zip_buffer.seek(0)
        
        return send_file(
            zip_buffer, 
            mimetype='application/zip', 
            download_name=f'Project_PIN_{data.get("egrknNumber", "draft")}.zip', 
            as_attachment=True
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return str(e), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)